import os
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import RGBColor, Pt
import re

TEMPLATES = {
    "PGD": {
        "path": "./templates/template_PGD.docx",
        "output_format": "docx"
    },

    "Thalassemia":{
        "path": "./templates/template_Thalassemia.docx",
        "output_format": "doc"
    }
}

MUTATION_PATTERNS = {
    r"(?i)\bSEA\b": "SEA",
    r"\b3\.7\b": "3.7",
    r"\b4\.2\b": "4.2",
    r"(?i)\bCD\s?-?\d+(-\d+)?\b": lambda m: re.sub(r"[\s\-]", "", m.group(0).upper()),  # e.g., CD 41-42 → CD41-42
    r"\bc\.\d+([+-]\d+)?[A-Z]?>[A-Z]?\b": lambda m: m.group(0),  # captures HGVS c. mutations
}

MUTATION_SEPARATOR_REGEX = r"\s*[,/&]\s*"

def filename_cleanup(s: str) -> str:
    s = re.sub(r'[<>:"/\\|?*]', '', s)
    s = s.replace('α', 'a').replace('Α', 'A')
    s = s.replace(' ', '_')
    s = s.replace('(', '').replace(')', '')
    return s

def extract_mutations(raw_text):
    if not isinstance(raw_text, str):
        return []

    results = []
    parts = re.split(MUTATION_SEPARATOR_REGEX, raw_text.strip())

    for part in parts:
        for pattern, normalizer in MUTATION_PATTERNS.items():
            match = re.search(pattern, part)
            if match:
                if callable(normalizer):
                    results.append(normalizer(match))
                elif normalizer:
                    results.append(normalizer)
                else:
                    results.append(match.group(0))
                break
    return results

def render_report(template_type, context, output_name, output_dir, embryos = None):
    config = TEMPLATES.get(template_type)
    if not config:
        raise FileNotFoundError(f"Không tìm thấy template cho {template_type}")
    
    template_path = config["path"]

    doc = DocxTemplate(template_path)
    doc.render(context)
    
    docx_path = os.path.join(output_dir, output_name + ".docx")
    doc.save(docx_path)

    return docx_path

def highlight_mutation_phrases(docx_path, phrases, color=(255, 0, 0), normal_color=(0, 112, 192)):
    doc = Document(docx_path)
    red = RGBColor(*color)
    blue = RGBColor(*normal_color)

    for para in doc.paragraphs:
        for phrase in phrases:
            if phrase.lower() in para.text.lower():
                full_text = para.text
                start = full_text.lower().find(phrase.lower())
                end = start + len(phrase)
                before = full_text[:start]
                match = full_text[start:end]
                after = full_text[end:]
                para.clear()

                before, match, after = full_text.partition(phrase)

                if before:
                    run_before = para.add_run(before)
                    run_before.font.color.rgb = blue
                    run_before.font.name = 'Times New Roman'
                    run_before.font.size = Pt(13)
                    run_before.font.bold = True

                run_match = para.add_run(match)
                run_match.font.color.rgb = red
                run_match.font.name = 'Times New Roman'
                run_match.font.size = Pt(13)
                run_match.font.bold = True

                if after:
                    run_after = para.add_run(after)
                    run_after.font.color.rgb = blue
                    run_after.font.name = 'Times New Roman'
                    run_after.font.size = Pt(13)
                    run_after.font.bold = True
                break

    doc.save(docx_path)

def extract_red_phrase(sentence: str) -> str:
    if "đột biến dị hợp tử" in sentence:
        start = sentence.find("đột biến dị hợp tử")
        end = sentence.find("trên gen", start)
        return sentence[start:end].strip()
    elif "dị hợp tử" in sentence:
        start = sentence.find("dị hợp tử")
        end = sentence.find("trên gen", start)
        return sentence[start:end].strip()
    return ""