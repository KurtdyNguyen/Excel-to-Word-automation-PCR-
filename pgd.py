import os
from datetime import datetime
import pandas as pd
from docx.shared import RGBColor, Pt
from docx import Document
from docxtpl import RichText
import re

from utils import (
    filename_cleanup,
    render_report,
    extract_mutations
)

def pgd_is_summary_row(row):
    return isinstance(row[4], str) and "VỢ" in row[4]

def pgd_parse_name_block(name_block):
    wife_name, wife_yob, husband_name, husband_yob = "", "", "", ""
    if isinstance(name_block, str):
        text = name_block.replace("CHỒNG:", "\nCHỒNG:").replace("VỢ:", "\nVỢ:")
        lines = [line.strip() for line in text.split("\n") if line.strip()]

        for line in lines:
            if "VỢ:" in line.upper():
                content = line.split(":", 1)[1].strip()
                parts = content.rsplit("-", 1) if "-" in content else content.rsplit(" ", 1)
                wife_name = parts[0].strip()
                wife_yob = parts[1].strip() if len(parts) > 1 else ""

            elif "CHỒNG:" in line.upper():
                content = line.split(":", 1)[1].strip()
                parts = content.rsplit("-", 1) if "-" in content else content.rsplit(" ", 1)
                husband_name = parts[0].strip()
                husband_yob = parts[1].strip() if len(parts) > 1 else ""

    return wife_name, wife_yob, husband_name, husband_yob

def style_embryo_table(docx_path):
    GREEN = RGBColor(0, 176, 80)
    RED = RGBColor(255, 0, 0)

    doc = Document(docx_path)
    for table in doc.tables:
        if len(table.rows) < 12:
            continue

        header = [cell.text.strip().lower() for cell in table.rows[10].cells]
        if "kết luận" not in "".join(header):
            continue

        for row in table.rows[11:]:
            cells = row.cells
            if len(cells) < 13:
                continue

            pgd_cell = cells[2]
            conclusion_cell = cells[10]

            concl_text = conclusion_cell.text.strip().lower()
            #if reading final conclusion paragraph, skip
            if any(keyword in concl_text for keyword in ["ghi chú", "kết luận", "hà nội"]):
                break
            if len(concl_text) > 50:
                break
            
            pgd_text = pgd_cell.text.strip().lower()

            if "đồng hợp" in pgd_text:
                color = RED
            elif "bất thường" in concl_text:
                color = RED
            elif "phôi có thể lựa chọn" in concl_text:
                color = GREEN
            else:
                color = None

            for cell in (pgd_cell, conclusion_cell):
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.bold = True
                        if color:
                            run.font.color.rgb = color

        break

    doc.save(docx_path)

def remove_alternating_empty_rows(docx_path):
    doc = Document(docx_path)

    def row_is_blank(row):
        return all(cell.text.strip().replace('\xa0', '').replace('\n', '').strip() == '' for cell in row.cells)


    for t_idx, table in enumerate(doc.tables):

        if len(table.rows) == 0:
            continue

        rows_to_remove = []
        for i in range(1, len(table.rows)):
            if i % 2 == 1 and row_is_blank(table.rows[i]):
                rows_to_remove.append(i)

        for i in reversed(rows_to_remove):
            table._tbl.remove(table.rows[i]._tr)

    doc.save(docx_path)

def generate_conclusion_block(homozygous, heterozygous, normal, mutation):
    paras = []

    if homozygous:
        rt = RichText()
        rt.add(f"Phôi {homozygous} có đồng hợp tử đột biến {mutation} gây bệnh nên không thể lựa chọn để cấy.", bold=True)
        paras.append(rt)

    if heterozygous:
        rt = RichText()
        rt.add(f"Phôi {heterozygous} có dị hợp tử đột biến {mutation} (phôi lành mang gen bệnh alpha thalassemia) và không phát hiện bất thường lệch bội số lượng nhiễm sắc thể nên có thể lựa chọn để cấy.", bold=True)
        paras.append(rt)

    if normal:
        rt = RichText()
        rt.add(f"Phôi {normal} không mang đột biến {mutation} gen bệnh alpha thalassemia nên có thể lựa chọn để cấy.", bold=True)
        paras.append(rt)

    return paras

def process_pgd_excel(file_path, output_dir):
    df = pd.read_excel(file_path, header=None)
    results = []

    i = 0
    while i < len(df):
        row = df.iloc[i]

        if pgd_is_summary_row(row):
            ID = filename_cleanup(str(row[5]) if pd.notna(row[5]) else "NO_ID")
            name_block = row[4]
            biopsy_date = row[1] if pd.notna(row[1]) else ""
            if isinstance(biopsy_date, pd.Timestamp):
                biopsy_date = biopsy_date.strftime("%d/%m/%Y")

            wife_name, wife_yob, husband_name, husband_yob = pgd_parse_name_block(name_block)

            # === Scan ahead to extract mutation info from embryo rows ===
            mutation = ""
            display_mutation = ""
            mutation_col = 8
            mutation_set = set()
            temp_i = i + 1
            while temp_i < len(df):
                next_row = df.iloc[temp_i]
                if pgd_is_summary_row(next_row) or pd.notna(next_row[0]):
                    break
                if isinstance(next_row[mutation_col], str):
                    found = extract_mutations(next_row[mutation_col])
                    mutation_set.update(found)
                temp_i += 1

            if mutation_set:
                display_mutation = ", ".join(mutation_set)
                mutation = display_mutation

            embryo_block_mutation = list(mutation_set)

            embryos = []
            i += 1
            while i < len(df):
                next_row = df.iloc[i]

                if pgd_is_summary_row(next_row) or pd.notna(next_row[0]):
                    break

                embryo_id = next_row[4]
                result = next_row[14]
                embryo_result_col = str(result).strip().lower() if pd.notna(result) else ""

                if "đồng hợp" in embryo_result_col:
                    zygosity = "Đồng hợp tử"
                elif "dị hợp" in embryo_result_col:
                    zygosity = "Dị hợp tử"
                elif "bình thường" in embryo_result_col:
                    zygosity = "Bình thường"
                else:
                    zygosity = ""

                embryo_mutations = extract_mutations(embryo_result_col)

                if not embryo_mutations and zygosity:
                    embryo_mutations = embryo_block_mutation

                if zygosity == "Bình thường":
                    pgd_result = "Bình thường"
                elif zygosity and embryo_mutations:
                    if len(embryo_mutations) == 1:
                        pgd_result = f"{zygosity} đột biến {embryo_mutations[0]}"
                    else:
                        mutation_label = " và đột biến ".join(embryo_mutations)
                        pgd_result = f"{zygosity} đột biến {mutation_label}"
                elif zygosity:
                    pgd_result = f"{zygosity} đột biến"
                else:
                    pgd_result = str(result).strip()

                if zygosity == "Bình thường":
                    final_conclusion = "Phôi có thể lựa chọn để cấy"
                elif zygosity == "Dị hợp tử":
                    if len(embryo_mutations) >= 2:
                        final_conclusion = "Bất thường"
                    else:
                        final_conclusion = "Phôi có thể lựa chọn để cấy"
                elif zygosity == "Đồng hợp tử":
                    final_conclusion = "Bất thường"
                else:
                    final_conclusion = ""

                embryos.append({
                    "index": len(embryos) + 1,
                    "name": str(embryo_id).strip() if pd.notna(embryo_id) else "",
                    "pgd_result": pgd_result,
                    "conclusion": final_conclusion
                })

                i += 1

            today = datetime.today()
            grouped = {
                "homozygous": ", ".join(e["name"] for e in embryos if "đồng hợp" in e["pgd_result"].lower()),
                "heterozygous": ", ".join(e["name"] for e in embryos if "dị hợp" in e["pgd_result"].lower()),
                "normal": ", ".join(e["name"] for e in embryos if "bình thường" in e["pgd_result"].lower())
            }
            context = {
                "ID": ID,
                "wife_name": wife_name,
                "wife_yob": wife_yob,
                "husband_name": husband_name,
                "husband_yob": husband_yob,
                "address": "",
                "mutation": mutation,
                "biopsy_date": biopsy_date,
                "date": str(today.day),
                "month": str(today.month),
                "year": str(today.year),
                **grouped,
                "conclusion_block": generate_conclusion_block(
                    grouped["homozygous"], grouped["heterozygous"], grouped["normal"], mutation
                ),
                "embryos": embryos
            }

            name_clean = filename_cleanup(wife_name)
            output_name = filename_cleanup(f"{ID}_{name_clean}_PGD")
            output_path = render_report("PGD", context, output_name, output_dir, embryos=embryos)
            style_embryo_table(output_path)
            remove_alternating_empty_rows(output_path)
            results.append((output_name, output_path))

        else:
            i += 1

    return results


if __name__ == "__main__":
    test_excel = "pgd.xlsx"
    output_dir = "test_output"
    os.makedirs(output_dir, exist_ok=True)

    results = process_pgd_excel(test_excel, output_dir)
    for output_name, output_path in results:
        print(f"Đã xuất file PGD: {output_name}")